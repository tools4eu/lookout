"""Report generator."""

import json
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape

from osint.core.config import Settings, get_settings, find_project_root
from osint.core.constants import IndicatorType, RiskLevel
from osint.models.reports import InvestigationReport, RelatedIndicator, TimelineEvent
from osint.orchestration.investigator import InvestigationResult


class ReportGenerator:
    """Generate reports from investigation results."""

    def __init__(self, settings: Optional[Settings] = None):
        """Initialize the report generator."""
        self.settings = settings or get_settings()

        # Set up Jinja2 environment
        template_dir = find_project_root() / "src" / "osint" / "reports" / "templates"
        if template_dir.exists():
            self.jinja_env = Environment(
                loader=FileSystemLoader(template_dir),
                autoescape=select_autoescape(["html", "xml"]),
            )
        else:
            self.jinja_env = None

    def create_report(
        self,
        result: InvestigationResult,
        analyst: Optional[str] = None,
    ) -> InvestigationReport:
        """
        Create an investigation report from results.

        Args:
            result: The investigation result
            analyst: Optional analyst name

        Returns:
            InvestigationReport
        """
        report = InvestigationReport(
            report_id=str(uuid.uuid4())[:8],
            analyst=analyst,
            indicator_value=result.indicator_value,
            indicator_type=result.indicator_type,
            risk_score=result.risk_score,
            risk_level=result.risk_level,
            sources_queried=[s.value for s in result.sources_queried],
            sources_with_data=[
                s.value for s, r in result.results.items() if r and r.success
            ],
        )

        # Generate risk summary
        report.risk_summary = report.generate_risk_summary()

        # Extract key findings
        report.key_findings = self._extract_key_findings(result)

        # Build timeline
        self._build_timeline(report, result)

        # Extract related indicators
        self._extract_related(report, result)

        # Generate executive summary
        report.executive_summary = self._generate_executive_summary(report)

        # Generate recommendations
        report.recommendations = self._generate_recommendations(report)

        return report

    def _extract_key_findings(self, result: InvestigationResult) -> list[str]:
        """Extract key findings from results."""
        from osint.models.results import (
            VirusTotalResult, URLScanResult, AbuseIPDBResult, ShodanResult,
            ThreatFoxResult, URLhausResult, CrtshResult,
        )

        findings = []

        for source, api_result in result.results.items():
            if not api_result or not api_result.success:
                continue

            # VirusTotal findings
            if isinstance(api_result, VirusTotalResult):
                mal = api_result.malicious or 0
                total = api_result.total_scanners or 0
                if mal > 0 and total > 0:
                    findings.append(
                        f"VirusTotal: {mal}/{total} scanners flagged as malicious"
                    )
                if api_result.categories:
                    cats = list(api_result.categories.values())[:2]
                    findings.append(f"VirusTotal categories: {', '.join(cats)}")

            # URLScan findings
            elif isinstance(api_result, URLScanResult):
                if api_result.malicious:
                    findings.append("URLScan: Flagged as malicious")
                if api_result.page_title and api_result.page_title != "404 Not Found":
                    findings.append(f"URLScan: Page title \"{api_result.page_title}\"")

            # AbuseIPDB findings
            elif isinstance(api_result, AbuseIPDBResult):
                score = api_result.abuse_confidence_score
                if score > 50:
                    findings.append(
                        f"AbuseIPDB: {score}% abuse confidence score"
                    )
                if api_result.total_reports and api_result.total_reports > 0:
                    findings.append(
                        f"AbuseIPDB: {api_result.total_reports} abuse reports"
                    )

            # ThreatFox findings
            elif isinstance(api_result, ThreatFoxResult):
                if api_result.total_matches > 0:
                    families = api_result.malware_families or []
                    if families:
                        findings.append(
                            f"ThreatFox: Associated with malware families: "
                            f"{', '.join(families[:3])}"
                        )
                    else:
                        findings.append(
                            f"ThreatFox: {api_result.total_matches} IOC matches found"
                        )

            # URLhaus findings
            elif isinstance(api_result, URLhausResult):
                if api_result.threat:
                    findings.append(
                        f"URLhaus: Classified as {api_result.threat}"
                    )

            # Shodan findings
            elif isinstance(api_result, ShodanResult):
                if api_result.vulns:
                    findings.append(
                        f"Shodan: {len(api_result.vulns)} known vulnerabilities"
                    )
                if api_result.ports:
                    findings.append(
                        f"Shodan: Open ports: {', '.join(map(str, api_result.ports[:5]))}"
                    )

            # crt.sh findings
            elif isinstance(api_result, CrtshResult):
                if api_result.subdomains and len(api_result.subdomains) > 5:
                    findings.append(
                        f"crt.sh: {len(api_result.subdomains)} subdomains found in certificates"
                    )

        # Add general finding if clean
        if not findings and result.risk_level == RiskLevel.CLEAN:
            findings.append("No malicious indicators detected across all sources")

        return findings

    def _build_timeline(
        self,
        report: InvestigationReport,
        result: InvestigationResult,
    ) -> None:
        """Build timeline from results."""
        from osint.models.results import (
            VirusTotalResult, RDAPResult, ThreatFoxResult, URLhausResult,
        )

        for source, api_result in result.results.items():
            if not api_result or not api_result.success:
                continue

            # RDAP/VT registration dates
            if isinstance(api_result, (RDAPResult, VirusTotalResult)):
                if api_result.creation_date:
                    report.add_timeline_event(
                        source=source.value,
                        event_type="registered",
                        description="Domain/Resource first registered",
                        timestamp=api_result.creation_date,
                    )

            # VirusTotal last analysis
            if isinstance(api_result, VirusTotalResult):
                if api_result.last_analysis_date:
                    report.add_timeline_event(
                        source=source.value,
                        event_type="scanned",
                        description=f"Last scanned by {source.value}",
                        timestamp=api_result.last_analysis_date,
                    )

            # ThreatFox dates
            if isinstance(api_result, ThreatFoxResult):
                if api_result.first_seen:
                    report.add_timeline_event(
                        source=source.value,
                        event_type="first_seen",
                        description="First seen in ThreatFox",
                        timestamp=api_result.first_seen,
                    )

            # URLhaus dates
            if isinstance(api_result, URLhausResult):
                if api_result.date_added:
                    report.add_timeline_event(
                        source=source.value,
                        event_type="added",
                        description="Added to URLhaus database",
                        timestamp=api_result.date_added,
                    )

    def _extract_related(
        self,
        report: InvestigationReport,
        result: InvestigationResult,
    ) -> None:
        """Extract related indicators from results."""
        from osint.models.results import (
            CrtshResult, ShodanResult, AlienVaultResult,
        )

        for source, api_result in result.results.items():
            if not api_result or not api_result.success:
                continue

            # Crt.sh subdomains
            if isinstance(api_result, CrtshResult):
                for subdomain in (api_result.subdomains or [])[:10]:
                    report.add_related_indicator(
                        value=subdomain,
                        indicator_type=IndicatorType.DOMAIN,
                        relationship="subdomain_of",
                        source=source.value,
                        confidence=0.9,
                    )

            # Shodan hostnames
            if isinstance(api_result, ShodanResult):
                for hostname in (api_result.hostnames or [])[:5]:
                    report.add_related_indicator(
                        value=hostname,
                        indicator_type=IndicatorType.DOMAIN,
                        relationship="resolves_to",
                        source=source.value,
                        confidence=0.8,
                    )

            # AlienVault related
            if isinstance(api_result, AlienVaultResult):
                for domain in (api_result.related_domains or [])[:5]:
                    report.add_related_indicator(
                        value=domain,
                        indicator_type=IndicatorType.DOMAIN,
                        relationship="related",
                        source=source.value,
                        confidence=0.6,
                    )

    def _generate_executive_summary(self, report: InvestigationReport) -> str:
        """Generate executive summary."""
        lines = [
            f"Investigation of {report.indicator_type.value} indicator: "
            f"{report.indicator_value}",
            "",
            report.risk_summary or "",
            "",
            f"Data was collected from {len(report.sources_with_data)} of "
            f"{len(report.sources_queried)} queried sources.",
        ]

        if report.key_findings:
            lines.append("")
            lines.append("Key findings include:")
            for finding in report.key_findings[:5]:
                lines.append(f"- {finding}")

        return "\n".join(lines)

    def _generate_recommendations(self, report: InvestigationReport) -> list[str]:
        """Generate recommendations based on risk level."""
        recs = []

        if report.risk_level == RiskLevel.CRITICAL:
            recs.extend([
                "Immediately block this indicator at network perimeter",
                "Search for historical connections in logs",
                "Conduct incident response if connections found",
                "Report to relevant threat intelligence sharing groups",
            ])
        elif report.risk_level == RiskLevel.HIGH:
            recs.extend([
                "Consider blocking this indicator",
                "Add to watchlist for monitoring",
                "Investigate any connections from internal systems",
            ])
        elif report.risk_level == RiskLevel.MEDIUM:
            recs.extend([
                "Add to monitoring watchlist",
                "Investigate if connections are from critical systems",
                "Consider additional sandboxing if file hash",
            ])
        elif report.risk_level == RiskLevel.LOW:
            recs.extend([
                "Continue monitoring",
                "No immediate action required",
            ])
        else:
            recs.extend([
                "Gather additional intelligence",
                "Monitor for future reports",
            ])

        return recs

    def to_markdown(self, report: InvestigationReport) -> str:
        """Render report as Markdown."""
        if self.jinja_env:
            try:
                template = self.jinja_env.get_template("report.md.j2")
                return template.render(report=report)
            except Exception:
                pass

        # Fallback to manual rendering
        lines = [
            f"# Investigation Report: {report.indicator_value}",
            "",
            f"**Report ID:** {report.report_id}",
            f"**Generated:** {report.created_at.isoformat()}",
            f"**Analyst:** {report.analyst or 'N/A'}",
            "",
            "## Executive Summary",
            "",
            report.executive_summary,
            "",
            "## Risk Assessment",
            "",
            f"- **Risk Level:** {report.risk_level.value.upper()}",
            f"- **Risk Score:** {report.risk_score:.0f}/100" if report.risk_score else "",
            "",
            report.risk_summary or "",
            "",
            "## Key Findings",
            "",
        ]

        for finding in report.key_findings:
            lines.append(f"- {finding}")

        lines.extend([
            "",
            "## Recommendations",
            "",
        ])

        for rec in report.recommendations:
            lines.append(f"- {rec}")

        if report.timeline:
            lines.extend([
                "",
                "## Timeline",
                "",
                "| Date | Source | Event | Description |",
                "|------|--------|-------|-------------|",
            ])
            for event in report.timeline:
                lines.append(
                    f"| {event.timestamp.strftime('%Y-%m-%d')} | "
                    f"{event.source} | {event.event_type} | {event.description} |"
                )

        if report.related_indicators:
            lines.extend([
                "",
                "## Related Indicators",
                "",
            ])
            for rel in report.related_indicators[:20]:
                lines.append(f"- **{rel.indicator_type.value}:** {rel.value} ({rel.relationship})")

        lines.extend([
            "",
            "---",
            "*Generated by OSINT Tool*",
        ])

        return "\n".join(lines)

    def to_json(self, report: InvestigationReport) -> str:
        """Render report as JSON."""
        return report.model_dump_json(indent=2)

    def to_docx(self, report: InvestigationReport) -> Document:
        """Render report as a Word document.

        Args:
            report: The investigation report to render.

        Returns:
            A python-docx Document object.
        """
        doc = Document()

        # ------------------------------------------------------------------ #
        # Page margins                                                         #
        # ------------------------------------------------------------------ #
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # ------------------------------------------------------------------ #
        # Helper: apply Calibri font to a paragraph run                       #
        # ------------------------------------------------------------------ #
        def _set_run_font(run, size_pt: int, bold: bool = False) -> None:
            run.font.name = "Calibri"
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            run.font.color.rgb = None  # inherit (black)

        def _heading(text: str, level_pt: int) -> None:
            """Add a heading paragraph styled in Calibri bold."""
            para = doc.add_paragraph()
            run = para.add_run(text)
            _set_run_font(run, level_pt, bold=True)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)

        def _body(text: str) -> None:
            """Add a normal body paragraph."""
            para = doc.add_paragraph()
            run = para.add_run(text)
            _set_run_font(run, 11)
            para.paragraph_format.space_after = Pt(4)

        # ------------------------------------------------------------------ #
        # 1. Title                                                             #
        # ------------------------------------------------------------------ #
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(
            f"Investigation Report: {report.indicator_value}"
        )
        _set_run_font(title_run, 16, bold=True)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(16)

        # ------------------------------------------------------------------ #
        # 2. Metadata table                                                    #
        # ------------------------------------------------------------------ #
        _heading("Report Details", 14)

        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = "Table Grid"
        meta_rows = [
            ("Report ID", report.report_id),
            ("Date", report.created_at.strftime("%Y-%m-%d %H:%M UTC")),
            ("Indicator Type", report.indicator_type.value),
            ("Risk Level", report.risk_level.value.upper()),
            (
                "Risk Score",
                f"{report.risk_score:.0f}/100" if report.risk_score is not None else "N/A",
            ),
        ]
        for i, (label, value) in enumerate(meta_rows):
            label_cell = meta_table.rows[i].cells[0]
            value_cell = meta_table.rows[i].cells[1]
            # Label (bold)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            _set_run_font(label_run, 11, bold=True)
            # Value
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            _set_run_font(value_run, 11)
        # Column widths
        for row in meta_table.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(11)

        doc.add_paragraph()  # spacer

        # ------------------------------------------------------------------ #
        # 3. Executive Summary                                                 #
        # ------------------------------------------------------------------ #
        _heading("Executive Summary", 14)
        if report.executive_summary:
            for line in report.executive_summary.splitlines():
                _body(line) if line.strip() else doc.add_paragraph()

        # ------------------------------------------------------------------ #
        # 4. Risk Assessment                                                   #
        # ------------------------------------------------------------------ #
        _heading("Risk Assessment", 14)
        risk_text = (
            f"{report.risk_level.value.upper()}"
            + (f" (score: {report.risk_score:.0f}/100)" if report.risk_score is not None else "")
        )
        _body(risk_text)
        if report.risk_summary:
            _body(report.risk_summary)

        # ------------------------------------------------------------------ #
        # 5. Key Findings                                                      #
        # ------------------------------------------------------------------ #
        _heading("Key Findings", 14)
        if report.key_findings:
            for finding in report.key_findings:
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(finding)
                _set_run_font(run, 11)
        else:
            _body("No key findings recorded.")

        # ------------------------------------------------------------------ #
        # 6. Recommendations                                                   #
        # ------------------------------------------------------------------ #
        _heading("Recommendations", 14)
        if report.recommendations:
            for idx, rec in enumerate(report.recommendations, start=1):
                para = doc.add_paragraph(style="List Number")
                run = para.add_run(rec)
                _set_run_font(run, 11)
        else:
            _body("No recommendations recorded.")

        # ------------------------------------------------------------------ #
        # 7. Timeline (only if populated)                                      #
        # ------------------------------------------------------------------ #
        if report.timeline:
            _heading("Timeline", 14)
            tl_headers = ["Date", "Source", "Event", "Description"]
            tl_table = doc.add_table(rows=1 + len(report.timeline), cols=4)
            tl_table.style = "Table Grid"
            # Header row
            hdr_cells = tl_table.rows[0].cells
            for col, header in enumerate(tl_headers):
                para = hdr_cells[col].paragraphs[0]
                run = para.add_run(header)
                _set_run_font(run, 11, bold=True)
            # Data rows
            for row_idx, event in enumerate(report.timeline, start=1):
                cells = tl_table.rows[row_idx].cells
                values = [
                    event.timestamp.strftime("%Y-%m-%d"),
                    event.source,
                    event.event_type,
                    event.description,
                ]
                for col, val in enumerate(values):
                    para = cells[col].paragraphs[0]
                    run = para.add_run(val)
                    _set_run_font(run, 11)
            # Column widths
            col_widths = [Cm(3), Cm(3), Cm(3.5), Cm(7)]
            for row in tl_table.rows:
                for col, width in enumerate(col_widths):
                    row.cells[col].width = width

            doc.add_paragraph()  # spacer

        # ------------------------------------------------------------------ #
        # 8. Related Indicators / Pivot Suggestions (only if populated)        #
        # ------------------------------------------------------------------ #
        if report.related_indicators:
            _heading("Related Indicators / Pivot Suggestions", 14)
            ri_headers = ["Type", "Value", "Relationship", "Source", "Confidence"]
            ri_table = doc.add_table(rows=1 + len(report.related_indicators[:20]), cols=5)
            ri_table.style = "Table Grid"
            # Header row
            hdr_cells = ri_table.rows[0].cells
            for col, header in enumerate(ri_headers):
                para = hdr_cells[col].paragraphs[0]
                run = para.add_run(header)
                _set_run_font(run, 11, bold=True)
            # Data rows
            for row_idx, rel in enumerate(report.related_indicators[:20], start=1):
                cells = ri_table.rows[row_idx].cells
                values = [
                    rel.indicator_type.value,
                    rel.value,
                    rel.relationship,
                    rel.source,
                    f"{rel.confidence:.0%}",
                ]
                for col, val in enumerate(values):
                    para = cells[col].paragraphs[0]
                    run = para.add_run(val)
                    _set_run_font(run, 11)
            # Column widths
            col_widths = [Cm(2.5), Cm(5), Cm(3.5), Cm(3), Cm(2.5)]
            for row in ri_table.rows:
                for col, width in enumerate(col_widths):
                    row.cells[col].width = width

            doc.add_paragraph()  # spacer

        # ------------------------------------------------------------------ #
        # 9. Data Sources                                                      #
        # ------------------------------------------------------------------ #
        _heading("Data Sources", 14)
        _body(f"Queried: {', '.join(report.sources_queried) if report.sources_queried else 'None'}")
        _body(
            f"With data: {', '.join(report.sources_with_data) if report.sources_with_data else 'None'}"
        )

        # ------------------------------------------------------------------ #
        # 10. Footer                                                           #
        # ------------------------------------------------------------------ #
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run("Generated by Lookout")
        footer_run.font.name = "Calibri"
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return doc

    def save_docx(self, report: InvestigationReport, output_path: Path) -> Path:
        """Save report as a Word document.

        Args:
            report: The investigation report to save.
            output_path: Destination path (extension added if absent).

        Returns:
            Path to the saved .docx file.
        """
        doc = self.to_docx(report)
        if not output_path.suffix:
            output_path = output_path.with_suffix(".docx")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def save_report(
        self,
        report: InvestigationReport,
        output_path: Path,
        format: str = "markdown",
    ) -> Path:
        """
        Save report to file.

        Args:
            report: The report to save
            output_path: Output file path
            format: Output format (markdown, json, docx)

        Returns:
            Path to saved file
        """
        if format == "json":
            content = self.to_json(report)
            if not output_path.suffix:
                output_path = output_path.with_suffix(".json")
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(content)
            return output_path
        elif format == "docx":
            return self.save_docx(report, output_path)
        else:
            content = self.to_markdown(report)
            if not output_path.suffix:
                output_path = output_path.with_suffix(".md")
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(content)
            return output_path
